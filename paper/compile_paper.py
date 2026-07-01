"""
compile_paper.py
Assembles all markdown sections into ParaKoop_Paper.docx.
Structure matches CES paper: no TOC, keywords after abstract,
notation + acknowledgements + declarations + references at end.
"""

import os, subprocess, re
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor

BASE = os.path.dirname(os.path.abspath(__file__))
FIGS = os.path.join(BASE, 'figures')

SECTIONS = [
    '00_frontmatter.md',
    '00_abstract.md',
    '01_introduction.md',
    '02_related_work.md',
    '03_methodology.md',
    '04_datasets.md',
    '05_experiments.md',
    '06_discussion.md',
    '07_conclusion.md',
    '09_endmatter.md',   # Notation + Acknowledgements + Declarations + References
]

FIGURE_CAPTIONS = {
    'fig1_architecture.png': (
        'Figure 1. ParaKoop architecture. '
        'Top row (cyan): forward flow θ → A(θ) → z* → [Cd, Cl]. '
        'Bottom row (amber): inverse design flow Cd* → AdamW gradient descent on θ → '
        'geometry changes in millimetres and degrees. '
        'The φ-grounding path (grey dashed) uses 75 AhmedML VTU flow fields to '
        'reduce the condition number κ(A(θ)) from 1.72–2.97 to 1.03–1.23.'
    ),
    'fig2_cd_scatter.png': (
        'Figure 2. Predicted versus CFD Cd for 74 held-out AhmedML cases (left) '
        'and prediction error distribution (right). '
        'Points are coloured by absolute error |ΔCd|. '
        'Mean |ΔCd| = 0.0378; median = 0.0364.'
    ),
    'fig3_hhl_analysis.png': (
        'Figure 3. HHL quantum interface analysis per body style. '
        '(a) Condition number κ(A(θ)); (b) theoretical HHL speedup over conjugate gradient '
        'using the formula K / (κ · log₂K) with K = 128; '
        '(c) predicted Cd at each body style mean geometry. '
        'All styles: 15 qubits, 100% solution fidelity.'
    ),
    'fig4_ablation.png': (
        'Figure 4. Proximity regularisation ablation (target Cd = 0.23, three body styles). '
        '(a) Height change from baseline; (b) width change from baseline; '
        '(c) |Cd_achieved − Cd_target|. '
        'Red bars: λ_prox = 0 (unconstrained, hits target exactly but produces implausible '
        'geometry); green bars: λ_prox = 2.0 (default, credible engineering suggestions).'
    ),
    'fig5_datasets.png': (
        'Figure 5. Training dataset composition. '
        '(a) Sample distribution by source (total ≈ 9,588 designs); '
        '(b) Cd distribution overlap between DrivAerNet++ STL (1,163 designs) and '
        'AhmedML held-out (74 cases), demonstrating that the unified θ '
        'bridges two datasets of different vehicle scales.'
    ),
    'fig6_eigenspectrum.png': (
        'Figure 6. Eigenvalue magnitudes |λ| of A(θ) for K = 128 modes, per body style. '
        'All modes cluster tightly near |λ| = 1.0 (vertical dashed line), '
        'confirming the near-identity operator design A(θ) = I + low-rank perturbation '
        'and explaining the low condition numbers κ = 1.03–1.23.'
    ),
    'fig7_streamlit_predict.png': (
        'Figure 7. ParaKoop Streamlit application — Predict tab. '
        'Geometry sliders (left panel) update the Cd/Cl speedometers in real time. '
        'Speedometer needle and value colours change with the performance zone '
        '(green: Cd < 0.25; amber: 0.25–0.35; red: > 0.35). '
        'Physics guardrail badges (Re, Ma, Eu) appear below the gauges.'
    ),
    'fig8_streamlit_design.png': (
        'Figure 8. ParaKoop Streamlit application — Design tab. '
        'User inputs a target Cd; the inverse design engine runs batch optimisation '
        'across three starting body styles and returns specific geometry changes '
        '(rear slant angle, height, width) in engineering units, '
        'with before/after comparison views and a Geometry Changes table.'
    ),
}

# Where to insert each figure (keyed on a unique string in the combined text)
FIGURE_INSERTIONS = [
    ('Figure 1 gives the full architecture overview', 'fig1_architecture.png', 'after_paragraph'),
    ('Figure 5 shows the training sample composition', 'fig5_datasets.png', 'before_paragraph'),
    ('Figure 6 plots the full eigenvalue magnitude', 'fig6_eigenspectrum.png', 'after_paragraph'),
    ('Figure 3 summarises κ, speedup', 'fig3_hhl_analysis.png', 'after_paragraph'),
    ('Figure 2 shows the full scatter', 'fig2_cd_scatter.png', 'after_paragraph'),
    ('Figure 4 compares height change', 'fig4_ablation.png', 'after_paragraph'),
    ('Figure 7', 'fig7_streamlit_predict.png', 'after_paragraph'),
    ('Figure 8', 'fig8_streamlit_design.png', 'after_paragraph'),
]


def strip_comments(text):
    return re.sub(r'<!--.*?-->', '', text, flags=re.DOTALL)


def strip_file_heading(text):
    """Remove '# Section N: ...' headings used only as file labels."""
    return re.sub(r'^# Section\s+\d+.*\n', '', text)


def image_md(fname):
    caption = FIGURE_CAPTIONS.get(fname, fname)
    img_path = os.path.join(FIGS, fname)
    return f'\n\n![{caption}]({img_path})\n\n'


def assemble():
    parts = []
    for fn in SECTIONS:
        path = os.path.join(BASE, fn)
        with open(path) as f:
            text = f.read()
        text = strip_comments(text)
        text = strip_file_heading(text)
        text = text.strip()

        # ── Insert figures in methodology ──────────────────────────────────────
        if fn == '03_methodology.md':
            text = text.replace(
                'Figure 1 gives the full architecture overview; the following subsections',
                'Figure 1 gives the full architecture overview; the following subsections'
            )
            # Insert figure block right after the line that mentions Figure 1
            text = re.sub(
                r'(Figure 1 gives the full architecture overview.*?detail\.)',
                r'\1' + image_md('fig1_architecture.png'),
                text, flags=re.DOTALL
            )

        # ── Insert figure in datasets ──────────────────────────────────────────
        if fn == '04_datasets.md':
            text = text.replace(
                '**Dataset statistics.** Figure 5 shows the training sample composition',
                image_md('fig5_datasets.png') +
                '**Dataset statistics.** Figure 5 shows the training sample composition'
            )

        # ── Insert figures in experiments ──────────────────────────────────────
        if fn == '05_experiments.md':
            # Fig 2 after §5.2 scatter mention (may wrap across lines)
            text = re.sub(
                r'(Figure 2\s+shows the full scatter.*?absolute error\.)',
                r'\1' + image_md('fig2_cd_scatter.png'),
                text, flags=re.DOTALL
            )
            # Fig 6 after §5.3 eigenspectrum mention
            text = text.replace(
                'Figure 6 plots the full eigenvalue magnitude distributions per body style.',
                'Figure 6 plots the full eigenvalue magnitude distributions per body style.' + image_md('fig6_eigenspectrum.png')
            )
            # Fig 3 after §5.4 qubit budget mention
            text = text.replace(
                'Figure 3 summarises κ, speedup, and predicted Cd per body style.',
                'Figure 3 summarises κ, speedup, and predicted Cd per body style.' + image_md('fig3_hhl_analysis.png')
            )
            # Fig 4 after §5.7 ablation mention
            text = text.replace(
                'Figure 4 compares height change, width change, and Cd error across both λ_prox settings\nfor all three body styles.',
                'Figure 4 compares height change, width change, and Cd error across both λ_prox settings\nfor all three body styles.' + image_md('fig4_ablation.png')
            )
            # Figs 7 & 8 in §5.8
            text = text.replace(
                'badges (Figure 7). The Design tab',
                'badges (Figure 7).' + image_md('fig7_streamlit_predict.png') + 'The Design tab'
            )
            text = text.replace(
                'Geometry Changes table (Figure 8).',
                'Geometry Changes table (Figure 8).' + image_md('fig8_streamlit_design.png')
            )

        # Strip any trailing --- so sections don't double-up when joined
        text = re.sub(r'\n\s*---\s*$', '', text.rstrip())
        parts.append(text)

    combined = '\n\n'.join(parts)
    combined_path = os.path.join(BASE, 'ParaKoop_combined.md')
    with open(combined_path, 'w') as f:
        f.write(combined)
    print(f'  Combined markdown: {combined_path}')
    return combined_path


def compile_word(combined_path):
    out = os.path.join(BASE, 'ParaKoop_Paper.docx')
    ref_doc = os.path.join(BASE, 'reference_template.docx')
    cmd = [
        'pandoc', combined_path,
        '--from', 'markdown+smart',
        '--to', 'docx',
        '--output', out,
        '--standalone',
        '--wrap=none',
        '--dpi=600',
        f'--reference-doc={ref_doc}',
    ]
    print('  Running pandoc ...')
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode == 0:
        size = os.path.getsize(out) / 1024
        print(f'  Word document: {out}  ({size:.0f} KB)')
    else:
        print('Pandoc error:')
        print(result.stderr)
    return result.returncode


INTUITION_TRIGGERS = ('Intuition —', 'What this section shows', 'Intuition—')
LIGHT_BLUE_FILL = 'D6EAF8'   # very light blue
BORDER_COLOR    = '2471A3'   # steel blue for borders


def _apply_para_box(para, fill_hex, border_color, border_sz=12):
    """Full-width box: fill + all-4-sides border + vertical spacing, zero indent."""
    pPr = para._p.get_or_add_pPr()

    # ── Background fill ──────────────────────────────────────────────────────
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill_hex)
    for e in pPr.findall(qn('w:shd')):
        pPr.remove(e)
    pPr.append(shd)

    # ── All-four-sides border ────────────────────────────────────────────────
    pBdr = OxmlElement('w:pBdr')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        # Left border is thicker as an accent bar
        sz = border_sz * 2 if edge == 'left' else border_sz
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    str(sz))
        el.set(qn('w:color'), border_color)
        el.set(qn('w:space'), '4')
        pBdr.append(el)
    for e in pPr.findall(qn('w:pBdr')):
        pPr.remove(e)
    pPr.append(pBdr)

    # ── Zero indent (full text-width, no extra indentation) ─────────────────
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'),  '0')
    ind.set(qn('w:right'), '0')
    for e in pPr.findall(qn('w:ind')):
        pPr.remove(e)
    pPr.append(ind)

    # ── Spacing: add breathing room above and below ──────────────────────────
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '80')
    spacing.set(qn('w:after'),  '80')
    for e in pPr.findall(qn('w:spacing')):
        pPr.remove(e)
    pPr.append(spacing)


def fix_frontmatter(docx_path):
    """Centre-align title and author block paragraphs."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = Document(docx_path)
    # Centre the first N paragraphs until we hit the Abstract heading
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading') and 'Abstract' in para.text:
            break
        if para.text.strip():
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.save(docx_path)
    print('  Frontmatter centred')


def add_page_numbers(docx_path):
    """Add page numbers to the footer of every section (centre-aligned)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document(docx_path)
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        # Clear existing footer content
        for para in footer.paragraphs:
            for run in para.runs:
                run.text = ''

        # Use first footer paragraph
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.clear()

        run = para.add_run()
        # Insert PAGE field
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = ' PAGE '
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

    doc.save(docx_path)
    print('  Page numbers added to footer')


def fix_intuition_boxes(docx_path):
    """Apply full-width light-blue box to all Intuition paragraphs."""
    doc = Document(docx_path)
    count = 0
    for para in doc.paragraphs:
        if any(para.text.startswith(t) for t in INTUITION_TRIGGERS):
            # Reset to Normal first so no style interferes with direct formatting
            para.style = doc.styles['Normal']
            _apply_para_box(para, LIGHT_BLUE_FILL, BORDER_COLOR, border_sz=12)
            count += 1
    doc.save(docx_path)
    print(f'  Intuition boxes styled: {count}')


def _set_cell_border(cell, **kwargs):
    """Apply border XML to a single table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        val = kwargs.get(edge, 'single')
        sz  = kwargs.get(f'{edge}_sz', 6)
        color = kwargs.get(f'{edge}_color', '000000')
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'),   val)
        el.set(qn('w:sz'),    str(sz))
        el.set(qn('w:color'), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def fix_tables(docx_path):
    """Bold all header rows and apply visible borders to every table."""
    doc = Document(docx_path)
    for table in doc.tables:
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                # Bold first row (header)
                if row_idx == 0:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.bold = True
                        # If cell has text but no runs, wrap it
                        if not para.runs and para.text.strip():
                            run = para.add_run(para.text)
                            run.bold = True
                            # Clear original text nodes
                            for child in list(para._p):
                                if child.tag != qn('w:r'):
                                    para._p.remove(child)
                # Borders on every cell
                _set_cell_border(cell,
                    top='single', top_sz=6, top_color='000000',
                    left='single', left_sz=6, left_color='000000',
                    bottom='single', bottom_sz=6, bottom_color='000000',
                    right='single', right_sz=6, right_color='000000',
                )
    doc.save(docx_path)
    print(f'  Tables fixed: {len(doc.tables)} tables styled')


if __name__ == '__main__':
    print('Compiling ParaKoop paper...')
    combined = assemble()
    rc = compile_word(combined)
    if rc == 0:
        out = os.path.join(BASE, 'ParaKoop_Paper.docx')
        fix_tables(out)
        fix_intuition_boxes(out)
        fix_frontmatter(out)
        add_page_numbers(out)
    print('Done.')
